import asyncio
import concurrent.futures
import html as html_lib
import io
import json
import re
import sys
import uuid
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Tuple

try:
    import docx
    from nicegui import app, background_tasks, context, events, ui

    import config
    from translate_core import (
        DocumentParser,
        Glossary,
        KnowledgeGraph,
        QAEngine,
        TranslationMemory,
        Translator,
    )
except ImportError as e:
    print(f"\n[ERROR] Import failed: {e.name}")
    sys.exit(1)

# Global resources
# ---------------------------------------------------------------------------

tm: "TranslationMemory | None" = None
glossary: "Glossary | None" = None
kg: "KnowledgeGraph | None" = None
translator: "Translator | None" = None
doc_parser: "DocumentParser | None" = None
qa_engine: "QAEngine | None" = None
llm_executor = concurrent.futures.ThreadPoolExecutor(max_workers=1)
GLOBAL_VOCAB: Dict[str, set] = {}  # Project ID -> Set of words

# ---------------------------------------------------------------------------
# Project persistence
# ---------------------------------------------------------------------------

PROJECTS_DIR = config.BASE_DIR / "data" / "projects"
PROJECTS_DIR.mkdir(parents=True, exist_ok=True)


def list_projects() -> list:
    projects = []
    for p in PROJECTS_DIR.glob("*.json"):
        try:
            data = json.loads(p.read_text(encoding="utf-8"))
            projects.append(
                {
                    "id": data["id"],
                    "filename": data["filename"],
                    "lang_pair": data["lang_pair"],
                    "saved_at": data["saved_at"],
                    "total": data["total"],
                    "done": data["done"],
                }
            )
        except Exception:
            pass
    projects.sort(key=lambda x: x["saved_at"], reverse=True)
    return projects


def save_project(ws: dict):
    segs = [
        {
            "id": s["id"],
            "source": s["source"],
            "target": s["target"],
            "status": s["status"],
        }
        for s in ws["segments"]
    ]
    done = sum(1 for s in segs if s["status"] == "done")
    data = {
        "id": ws["project_id"],
        "filename": ws["filename"],
        "lang_pair": ws["lang_pair"],
        "active_index": ws["active_index"],
        "saved_at": datetime.now().isoformat(timespec="seconds"),
        "total": len(segs),
        "done": done,
        "segments": segs,
    }
    path = PROJECTS_DIR / f"{ws['project_id']}.json"
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def load_project(project_id: str):
    path = PROJECTS_DIR / f"{project_id}.json"
    if not path.exists():
        return None
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return None


def delete_project(project_id: str):
    for suffix in [".json", ".docx", ".txt"]:
        p = PROJECTS_DIR / f"{project_id}{suffix}"
        if p.exists():
            p.unlink()


# ---------------------------------------------------------------------------
# Working TM
# ---------------------------------------------------------------------------


def save_pair_to_tm(source: str, target: str, lang_pair: str):
    src_lang, tgt_lang = lang_pair.split("->")
    tm_path = config.TM_DIR / "working.tmx"
    config.TM_DIR.mkdir(parents=True, exist_ok=True)

    if not tm_path.exists():
        header = (
            "\n".join(
                [
                    '<?xml version="1.0" encoding="UTF-8"?>',
                    '<tmx version="1.4">',
                    f'  <header creationtool="ZenTranslator" srclang="{src_lang}"/>',
                    "  <body>",
                    "  </body>",
                    "</tmx>",
                ]
            )
            + "\n"
        )
        tm_path.write_text(header, encoding="utf-8")

    def _esc(s: str) -> str:
        return (
            s.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
        )

    tu = (
        "\n".join(
            [
                "    <tu>",
                f'      <tuv xml:lang="{src_lang}"><seg>{_esc(source)}</seg></tuv>',
                f'      <tuv xml:lang="{tgt_lang}"><seg>{_esc(target)}</seg></tuv>',
                "    </tu>",
            ]
        )
        + "\n"
    )

    raw = tm_path.read_text(encoding="utf-8")
    raw = raw.replace("  </body>", tu + "  </body>")
    tm_path.write_text(raw, encoding="utf-8")

    if tm is not None:
        tm.entries.append({"source": source, "target": target, "origin": "working.tmx"})


class SmartPhraseExtractor:
    def get_ngrams(self, text: str, n: int = 3):
        words = re.findall(r"\w+[\w'-]*\w+", text.lower())
        return [" ".join(words[i : i + n]) for i in range(len(words) - n + 1)]

    def scan_for_phrases(self, source_text: str):
        cands: set = set()
        for n in [4, 3, 2]:
            cands.update(self.get_ngrams(source_text, n))
        cands.update(
            w for w in re.findall(r"\w+[\w'-]*\w+", source_text.lower()) if len(w) > 3
        )
        return list(cands)


extractor = SmartPhraseExtractor()


@app.on_startup
async def init_resources():
    global tm, glossary, kg, translator, doc_parser, qa_engine
    tm = TranslationMemory()
    glossary = Glossary()
    kg = KnowledgeGraph()
    doc_parser = DocumentParser()
    translator = Translator()
    qa_engine = QAEngine()

    loop = asyncio.get_running_loop()
    loop.run_in_executor(llm_executor, lambda: translator._impl._ensure_loaded())


# ===========================================================================
# PAGE /  —  Project list
# ===========================================================================


@ui.page("/")
def page_home():
    _apply_colors()
    ui.add_head_html("""<style>
        .proj-card {
            transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
            cursor: pointer;
            border: 1px solid #e2e8f0;
        }
        .proj-card:hover {
            box-shadow: 0 20px 25px -5px rgba(0, 0, 0, 0.05), 0 10px 10px -5px rgba(0, 0, 0, 0.02);
            transform: translateY(-4px);
            border-color: #3b82f6;
        }
        .hero-gradient {
            background: linear-gradient(135deg, #0f172a 0%, #1e293b 100%);
        }
    </style>""")

    with ui.column().classes("w-full min-h-screen bg-slate-50"):
        with ui.row().classes(
            "w-full hero-gradient px-8 py-12 items-center justify-between shadow-lg"
        ):
            with ui.row().classes("items-center gap-4"):
                ui.icon("blur_on", size="56px").classes("text-blue-400 animate-pulse")
                with ui.column().classes("gap-0"):
                    ui.label("Zen Translator").classes(
                        "text-3xl font-black text-white tracking-tighter"
                    )
                    ui.label("Professional Translation Workspace").classes(
                        "text-blue-200/60 text-xs font-bold uppercase tracking-[0.2em]"
                    )

            with ui.row().classes("gap-4"):
                ui.label("V 2.0").classes(
                    "text-white/20 text-[10px] font-black border border-white/10 px-2 py-1 rounded"
                )

        with ui.column().classes("w-full max-w-4xl mx-auto px-6 -mt-8 gap-8 pb-20"):
            with ui.card().classes(
                "w-full p-8 rounded-3xl border-none bg-white shadow-2xl"
            ):
                with ui.row().classes("w-full items-center justify-between mb-6"):
                    with ui.column().classes("gap-1"):
                        ui.label("Start a New Translation").classes(
                            "text-xl font-bold text-slate-800"
                        )
                        ui.label("Upload a .docx file to begin your project").classes(
                            "text-sm text-slate-400"
                        )

                    with ui.row().classes(
                        "gap-3 items-center bg-slate-50 p-2 rounded-2xl border border-slate-100"
                    ):
                        lang_opts = ["en", "sl", "de", "fr", "it"]
                        src_lang = (
                            ui.select(lang_opts, value="en")
                            .props("outlined dense rounded")
                            .classes("w-20 bg-white")
                        )
                        ui.icon("swap_horiz", size="sm").classes("text-slate-300")
                        tgt_lang = (
                            ui.select(lang_opts, value="sl")
                            .props("outlined dense rounded")
                            .classes("w-20 bg-white")
                        )

                async def upload_wrapper(e):
                    await _handle_new_upload(e, f"{src_lang.value}->{tgt_lang.value}")

                ui.upload(
                    on_upload=upload_wrapper,
                    auto_upload=True,
                    label="Drop files here or click to browse",
                    max_files=1,
                ).classes("w-full").props("color=accent accept=.docx flat bordered")

            proj_container = ui.column().classes("w-full gap-4")
            _render_project_list(proj_container)


def _render_project_list(container: ui.column):
    container.clear()
    projects = list_projects()
    with container:
        if not projects:
            with ui.column().classes("w-full items-center py-20 opacity-20"):
                ui.icon("folder_open", size="64px")
                ui.label("No active projects").classes("text-lg font-bold")
            return

        ui.label("RECENT PROJECTS").classes(
            "text-[10px] font-black text-slate-400 uppercase tracking-[0.3em] px-2 mb-2"
        )
        with ui.grid(columns=2).classes("w-full gap-4"):
            for p in projects:
                pct = int(p["done"] / p["total"] * 100) if p["total"] else 0
                with (
                    ui.card()
                    .classes("p-6 rounded-2xl bg-white proj-card flex flex-col gap-4")
                    .on(
                        "click", lambda pid=p["id"]: ui.navigate.to(f"/translate/{pid}")
                    )
                ):
                    with ui.row().classes("w-full justify-between items-start"):
                        with ui.column().classes("gap-0.5 flex-1"):
                            ui.label(p["filename"]).classes(
                                "text-base font-bold text-slate-800 truncate leading-tight"
                            )
                            ui.label(p["lang_pair"]).classes(
                                "text-[10px] font-black text-blue-500 uppercase tracking-widest"
                            )
                        ui.button(
                            icon="delete",
                            on_click=lambda e, pid=p["id"], c=container: (
                                _delete_and_refresh(pid, c)
                            ),
                        ).props("flat round dense size=sm color=slate-300").classes(
                            "hover:text-red-400 transition-colors"
                        ).on("click.stop")

                    with ui.row().classes("w-full items-center gap-4 mt-2"):
                        with ui.column().classes("flex-1 gap-1"):
                            ui.linear_progress(value=pct / 100, color="positive").props(
                                "size=8px rounded"
                            ).classes("w-full")
                            with ui.row().classes(
                                "w-full justify-between items-center"
                            ):
                                ui.label(
                                    f"{p['done']} / {p['total']} segments"
                                ).classes("text-[10px] text-slate-400 font-medium")
                                ui.label(f"{pct}%").classes(
                                    "text-[10px] font-bold text-slate-600"
                                )

                    with ui.row().classes(
                        "w-full border-t border-slate-50 pt-4 mt-auto items-center justify-between"
                    ):
                        ui.label(
                            f"Saved {p['saved_at'][:16].replace('T', ' ')}"
                        ).classes("text-[9px] text-slate-300 font-medium italic")
                        ui.icon("arrow_forward", size="14px").classes("text-blue-200")


def _delete_and_refresh(project_id: str, container: ui.column):
    cli = context.client
    with cli:
        ui.notify("Deleted", type="warning", timeout=1200)
        delete_project(project_id)
        _render_project_list(container)


async def _handle_new_upload(e, lang_pair: str):
    name = getattr(e, "name", "document.docx")
    if not name.endswith(".docx"):
        return ui.notify(".docx only", type="warning")
    try:
        content = await e.file.read()
        doc = docx.Document(io.BytesIO(content))
    except Exception as ex:
        return ui.notify(f"Error: {ex}", type="negative")

    project_id = str(uuid.uuid4())[:8]
    segments = []
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            segments.append(
                {"id": len(segments), "source": txt, "target": "", "status": "pending"}
            )

    ws = {
        "project_id": project_id,
        "filename": name,
        "lang_pair": lang_pair,
        "active_index": 0,
        "segments": segments,
    }
    (PROJECTS_DIR / f"{project_id}.docx").write_bytes(content)
    save_project(ws)
    ui.notify(f"Created: {len(segments)} segments", type="positive")
    ui.navigate.to(f"/translate/{project_id}")


def _get_highlighted_source(text: str) -> str:
    if kg is None:
        return text
    entities = kg.extract_entities(text)
    if not entities:
        return text

    html = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    sorted_entities = sorted(
        entities, key=lambda x: len(x.get("term", "")), reverse=True
    )

    for e in sorted_entities:
        term = e.get("term")
        if not term:
            continue
        pattern = re.compile(f"({re.escape(term)})", re.IGNORECASE)
        html = pattern.sub(r'<span class="kg-term">\1</span>', html)

    return html


def _run_qa(seg: dict, lang_pair: str) -> List[Dict]:
    if qa_engine is None:
        return []
    src, tgt = lang_pair.split("->")
    g_hits = glossary.lookup_terms(seg["source"], src, tgt) if glossary else []
    return qa_engine.check_segment(seg["source"], seg["target"], g_hits)


def _search_intelligence(query: str, lang_pair: str) -> Dict:
    results = {"kg": [], "tm": []}
    if not query or len(query) < 2:
        return results

    src_l, tgt_l = lang_pair.split("->")

    if kg:
        for node_id, data in kg.G.nodes(data=True):
            term = (
                data.get("term") or data.get("label") or data.get("id") or str(node_id)
            )
            if term and isinstance(term, str) and query.lower() in term.lower():
                rels = kg.find_neighbors(node_id, max_depth=1)
                results["kg"].append({**data, "relations": rels})

    if tm:
        hits = tm.search_concordance(query, top_n=10)
        results["tm"] = hits

    return results


# ===========================================================================
# PAGE /translate/{project_id}  —  Translation workspace
# ===========================================================================


@ui.page("/translate/{project_id}")
def page_translate(project_id: str):
    _apply_colors()

    # Styles for this page
    ui.add_head_html("""<style>
        .intel-header {
            font-size: 9px;
            font-weight: 900;
            color: #94a3b8;
            letter-spacing: 0.2em;
            text-transform: uppercase;
        }
        .source-area {
            background: #f8fafc;
            border-radius: 12px;
            padding: 16px;
            border: 1px solid #e2e8f0;
        }
        .active-card {
            box-shadow: 0 0 0 2px #3b82f6, 0 20px 40px -8px rgba(59,130,246,0.15);
        }
        .kg-term {
            background: #eff6ff;
            color: #1d4ed8;
            border-radius: 3px;
            padding: 0 3px;
            font-weight: 600;
        }
        .inline-ghost {
            position: absolute;
            pointer-events: none;
            color: #9ca3af;
            opacity: 0.7;
            white-space: pre;
            z-index: 10;
        }
    </style>""")

    data = load_project(project_id)
    if data is None:
        ui.label("Project not found.").classes("m-8 text-red-500")
        ui.button("← Back", on_click=lambda: ui.navigate.to("/")).classes("m-4")
        return

    # --- Intelligence Panel Logic ---
    def _add_to_kg(src: str, tgt: str, notify: bool = True):
        src_l, tgt_l = data["lang_pair"].split("->")
        kg.promote_pair(src, tgt, src_l, tgt_l, verified=True)
        kg.save()
        if notify:
            ui.notify(
                f"Promoted '{src}' to Knowledge Graph",
                type="positive",
                icon="auto_awesome",
            )

    def _scan_tm_for_kg():
        if not tm:
            return
        count = 0
        try:
            for entry in tm.entries[:50]:
                if not any(entry["source"].lower() in n.lower() for n in kg.G.nodes):
                    if 1 < len(entry["source"].split()) < 3:
                        _add_to_kg(entry["source"], entry["target"], notify=False)
                        count += 1
            print(f"[KG Scan] Discovery complete: Added {count} new entities.")
        except Exception as e:
            print(f"[KG Scan] Error: {e}")

    @ui.refreshable
    def _render_intel(query: str = ""):
        res = _search_intelligence(query, data["lang_pair"])
        if not query:
            ui.label(
                "Search for terms to see Knowledge Graph relations and TM concordance hits."
            ).classes("text-slate-400 italic text-xs mt-4 leading-relaxed")
            with ui.column().classes(
                "w-full mt-8 pt-6 border-t border-slate-100 gap-4"
            ):
                ui.label("SMART DISCOVERY").classes(
                    "text-[9px] font-black text-slate-400 tracking-[.2em]"
                )
                ui.button(
                    "Scan TM for Entities", icon="psychology", on_click=_scan_tm_for_kg
                ).classes(
                    "w-full bg-white text-slate-600 border border-slate-200 rounded-xl py-4 shadow-sm hover:bg-slate-50 transition-all no-wrap"
                ).props("flat")
            return

        if res["kg"]:
            with ui.column().classes("w-full gap-2 mt-4"):
                ui.label("KNOWLEDGE GRAPH").classes(
                    "text-[9px] font-black text-blue-500 tracking-[.2em]"
                )
                for node in res["kg"][:5]:
                    with ui.card().classes(
                        "w-full p-3 bg-blue-50/40 border-blue-100 rounded-xl shadow-sm"
                    ):
                        label = node.get("term") or node.get("label") or node.get("id")
                        with ui.row().classes("items-center gap-2"):
                            ui.icon("account_tree", size="14px", color="blue-500")
                            ui.label(label).classes("font-bold text-slate-800 text-sm")
                        if node.get("relations"):
                            ui.separator().classes("my-2 opacity-50")
                            for r in node.get("relations", []):
                                r_label = r.get("term") or r.get("label") or r.get("id")
                                ui.label(f"• {r.get('relation')}: {r_label}").classes(
                                    "text-[11px] text-slate-600 ml-1"
                                )

        if res["tm"]:
            with ui.column().classes("w-full gap-2 mt-6"):
                ui.label("TM CONCORDANCE").classes(
                    "text-[9px] font-black text-slate-400 tracking-[.2em]"
                )
                src_l, tgt_l = data["lang_pair"].split("->")
                highlight_terms = {query.lower()}

                if glossary:
                    for m in glossary.search_prefix(query, src_l, tgt_l):
                        highlight_terms.add(m.lower())
                    direct_g = glossary.lookup_terms(query, src_l, tgt_l)
                    for g in direct_g:
                        highlight_terms.add(g["source_term"].lower())
                        highlight_terms.add(g["target_term"].lower())

                if kg:
                    for node_id, d in kg.G.nodes(data=True):
                        term = d.get("term") or d.get("label")
                        if term and (
                            term.lower().startswith(query.lower())
                            or query.lower() in term.lower()
                        ):
                            highlight_terms.add(term.lower())
                            for r in kg.find_neighbors(node_id, max_depth=2):
                                if r.get("type") == "term" and r.get("term"):
                                    highlight_terms.add(r.get("term").lower())

                valid_terms = sorted(
                    [t for t in highlight_terms if len(t) >= 2], key=len, reverse=True
                )
                pattern = (
                    re.compile(
                        f"({'|'.join(re.escape(t) for t in valid_terms)})",
                        re.IGNORECASE,
                    )
                    if valid_terms
                    else None
                )

                for hit in res["tm"][:8]:
                    with ui.card().classes(
                        "w-full p-4 bg-white border-slate-100 rounded-2xl shadow-sm hover:shadow-md hover:border-blue-200 transition-all cursor-pointer group"
                    ):
                        with ui.row().classes(
                            "w-full justify-between items-start no-wrap"
                        ):
                            with ui.column().classes("flex-1"):
                                if pattern:
                                    marked_src = pattern.sub(
                                        r'<mark class="bg-amber-100 text-amber-900 rounded px-1 font-medium">\1</mark>',
                                        hit["source"],
                                    )
                                    marked_tgt = pattern.sub(
                                        r'<mark class="bg-amber-100 text-amber-900 rounded px-1 font-medium">\1</mark>',
                                        hit["target"],
                                    )
                                else:
                                    marked_src, marked_tgt = (
                                        hit["source"],
                                        hit["target"],
                                    )
                                ui.html(marked_src).classes(
                                    "text-[11px] text-slate-400 leading-snug mb-1.5"
                                )
                                ui.html(marked_tgt).classes(
                                    "text-sm text-slate-800 font-medium leading-relaxed"
                                )
                            ui.button(
                                icon="add_circle",
                                on_click=lambda h=hit: _add_to_kg(
                                    h["source"], h["target"]
                                ),
                            ).props("flat round dense size=sm color=blue-400").classes(
                                "opacity-0 group-hover:opacity-100 transition-opacity ml-2"
                            ).tooltip("Promote to KG")

        if not res["kg"] and not res["tm"]:
            ui.label("No matches found in Intelligence databases.").classes(
                "text-slate-300 italic text-xs mt-4"
            )

    # --- Layout ---
    with (
        ui.right_drawer(value=True, fixed=True)
        .classes("bg-slate-50 border-l border-slate-100 p-6 shadow-2xl")
        .props("width=380") as intel_drawer
    ):
        with ui.row().classes("w-full justify-between items-center mb-6"):
            ui.label("INTELLIGENCE").classes(
                "text-[11px] font-black text-slate-500 tracking-[.3em]"
            )
            ui.button(icon="close", on_click=intel_drawer.toggle).props(
                "flat round dense size=sm color=slate-300"
            )
        with ui.column().classes("w-full gap-4"):
            search_field = (
                ui.input(
                    placeholder="Instant Search (TM/KG)...",
                    on_change=lambda e: _render_intel.refresh(e.value),
                )
                .props("outlined dense clearable bg-white rounded-xl")
                .classes("w-full shadow-sm")
            )
            with ui.scroll_area().classes("w-full h-[calc(100vh-200px)]"):
                _render_intel()

    ws = {
        "project_id": project_id,
        "filename": data["filename"],
        "lang_pair": data["lang_pair"],
        "active_index": data.get("active_index", 0),
        "segments": data["segments"],
        "is_batch": False,
    }

    def _progress() -> float:
        if not ws["segments"]:
            return 0.0
        return sum(1 for s in ws["segments"] if s["status"] == "done") / len(
            ws["segments"]
        )

    def _autosave():
        save_project(ws)

    def _set_active(idx: int):
        if 0 <= idx < len(ws["segments"]):
            cli = context.client
            ws["active_index"] = idx
            _autosave()
            with cli:
                _flow.refresh()
                card_id = ws["segments"][idx].get("_card_id")
                if card_id and cli.has_socket_connection:
                    try:
                        ui.run_javascript(
                            f'setTimeout(() => {{ const el = getHtmlElement({card_id}); if(el) el.scrollIntoView({{behavior: "smooth", block: "center"}}); }}, 50);'
                        )
                    except KeyError as e:
                        if "Session is disconnected" in str(e):
                            return
                        raise

    async def _confirm_segment():
        idx = ws["active_index"]
        seg = ws["segments"][idx]

        if not seg["target"].strip():
            return

        seg["status"] = "done"
        _add_to_kg(seg["source"], seg["target"])
        save_pair_to_tm(seg["source"], seg["target"], ws["lang_pair"])

        next_idx = idx + 1
        if next_idx < len(ws["segments"]):
            _set_active(next_idx)
            with context.client:
                _bar.refresh()
        else:
            ws["active_index"] = next_idx
            with context.client:
                ui.notify("Document complete! 🎉", type="positive")
                _bar.refresh()
                _flow.refresh()

    def _open_glossary():
        with ui.dialog() as dialog, ui.card().classes("min-w-[400px]"):
            ui.label("Add to Glossary").classes("text-lg font-bold mb-2")
            src_lang, tgt_lang = ws["lang_pair"].split("->")
            src_input = ui.input(f"Source Term ({src_lang})").classes("w-full")
            tgt_input = ui.input(f"Target Term ({tgt_lang})").classes("w-full")
            note_input = ui.input("Note (optional)").classes("w-full")

            def save():
                s = src_input.value.strip() if src_input.value else ""
                t = tgt_input.value.strip() if tgt_input.value else ""
                if s and t:
                    entry = f"{s}\t{t}\t{note_input.value or ''}\n"
                    custom_path = config.GLOSSARY_DIR / "custom.tsv"
                    config.GLOSSARY_DIR.mkdir(parents=True, exist_ok=True)
                    with open(custom_path, "a", encoding="utf-8") as f:
                        f.write(entry)
                    if glossary:
                        glossary._add_simple_entry(
                            s, t, src_lang, tgt_lang, "custom.tsv", note_input.value
                        )
                        glossary._build_indices()
                    ui.notify("Term added to glossary", type="positive")
                    dialog.close()
                    _flow.refresh()
                else:
                    ui.notify("Both terms are required", type="negative")

            with ui.row().classes("w-full justify-end mt-4"):
                ui.button("Cancel", on_click=dialog.close).props("flat")
                ui.button("Save", on_click=save).props("color=positive")
        dialog.open()

    async def _ai(seg: dict, ti: ui.textarea, force: bool = False):
        if not force and seg["target"].strip():
            return
        src, tgt = ws["lang_pair"].split("->")
        loop = asyncio.get_running_loop()

        def _look():
            a = tm.lookup_fuzzy(seg["source"], threshold=90.0, limit=1) if tm else []
            b = glossary.lookup_terms(seg["source"], src, tgt) if glossary else []
            c = tm.search_concordance(seg["source"], top_n=2) if tm else []
            k = kg.extract_entities(seg["source"]) if kg else []
            return a, b, c, k

        try:
            ti.props("loading")
            orig = seg["target"]
        except RuntimeError:
            return

        try:
            a, b, c, k = await loop.run_in_executor(None, _look)
            _, text = await loop.run_in_executor(
                llm_executor,
                lambda: translator.translate(seg["source"], src, tgt, a, b, c, k),
            )
            try:
                if seg["target"] == orig:
                    seg["target"] = text
                    seg["_ai_draft"] = text
                    if "_sugg_cache" in seg:
                        del seg["_sugg_cache"]
                    _autosave()
            except RuntimeError:
                pass
        except Exception as ex:
            print(f"[AI] {ex}")
            ui.notify("AI failed", type="negative")
        finally:
            try:
                ti.props(remove="loading")
            except RuntimeError:
                pass

    async def _inline_panel(seg: dict, ti: ui.textarea, container: ui.column):
        _panel_client = container.client
        src, tgt = ws["lang_pair"].split("->")
        loop = asyncio.get_running_loop()

        def _look():
            a = tm.lookup_fuzzy(seg["source"], threshold=90.0, limit=3) if tm else []
            b = glossary.lookup_terms(seg["source"], src, tgt) if glossary else []
            c = tm.search_concordance(seg["source"], top_n=3) if tm else []
            return a, b, c

        tm_m, g_h, c_h = await loop.run_in_executor(None, _look)

        print(
            f"[Inline] _inline_panel results - TM: {len(tm_m)}, Glossary: {len(g_h)}, Concordance: {len(c_h)}"
        )
        try:
            if container.is_deleted or not _panel_client.has_socket_connection:
                return
            with _panel_client:
                container.clear()
        except RuntimeError:
            return

        def _ins(txt: str):
            try:
                cur = seg["target"] or ""
                sep = " " if cur and not cur.endswith(" ") else ""
                nv = (cur + sep + txt).strip()
                seg["target"] = nv
            except RuntimeError:
                pass

        def _rep(txt: str):
            try:
                seg["target"] = txt
            except RuntimeError:
                pass

        try:
            with _panel_client:
                with container:
                    has = False
                    if g_h:
                        has = True
                        with ui.column().classes("w-full gap-2 mb-2"):
                            ui.label("GLOSSARY HITS").classes("intel-header")
                            with ui.row().classes(
                                "w-full gap-2 items-center flex-wrap"
                            ):
                                for g in g_h:
                                    ui.button(
                                        f"{g['target_term']}",
                                        on_click=lambda e, g=g: _ins(g["target_term"]),
                                    ).props("unelevated rounded").classes(
                                        "bg-emerald-100 hover:bg-emerald-200 text-[11px] font-bold text-emerald-800 px-3 h-8 normal-case"
                                    ).tooltip(
                                        f"{g['source_term']} → {g['target_term']} | {g.get('note') or 'Click to insert'}"
                                    )
                    if tm_m:
                        has = True
                        with ui.column().classes("w-full gap-2 mt-2"):
                            ui.label("TM MATCHES").classes("intel-header")
                            for m in tm_m:
                                score = int(m.get("score", 0))
                                badge = (
                                    "text-emerald-700 bg-emerald-50 border-emerald-200"
                                    if score >= 85
                                    else "text-blue-700 bg-blue-50 border-blue-200"
                                )
                                with (
                                    ui.row()
                                    .classes(
                                        "w-full bg-white border border-slate-200 rounded-xl p-3 items-center gap-4 hover:border-blue-300 transition-colors cursor-pointer"
                                    )
                                    .on("click", lambda e, m=m: _rep(m["target"]))
                                ):
                                    ui.label(f"{score}%").classes(
                                        f"text-[10px] font-black px-2 py-1 rounded-lg border {badge} shrink-0"
                                    )
                                    with ui.column().classes("gap-0.5 flex-1 min-w-0"):
                                        ui.label(m["source"]).classes(
                                            "text-[11px] text-slate-400 italic leading-snug"
                                        ).style(
                                            "white-space:normal;word-break:break-word"
                                        )
                                        ui.label(m["target"]).classes(
                                            "text-[13px] text-slate-800 font-bold leading-snug"
                                        ).style(
                                            "white-space:normal;word-break:break-word"
                                        )
                                    ui.icon("content_paste", size="18px").classes(
                                        "text-slate-300"
                                    )
                    srcs = {m["source"] for m in tm_m}
                    conc = [c for c in c_h if c["source"] not in srcs]
                    if conc:
                        has = True
                        with ui.column().classes("w-full gap-2 mt-2"):
                            ui.label("CONCORDANCE").classes("intel-header")
                            for c in conc[:2]:
                                with (
                                    ui.row()
                                    .classes(
                                        "w-full bg-white/50 border border-slate-100 rounded-xl p-3 items-start gap-4 hover:border-slate-300 transition-colors cursor-pointer"
                                    )
                                    .on("click", lambda e, c=c: _ins(c["target"]))
                                ):
                                    ui.icon("search", size="16px").classes(
                                        "text-slate-300 mt-1 shrink-0"
                                    )
                                    with ui.column().classes("gap-0.5 flex-1"):
                                        ui.label(c["source"]).classes(
                                            "text-[11px] text-slate-400 italic leading-snug"
                                        ).style(
                                            "white-space:normal;word-break:break-word"
                                        )
                                        ui.label(c["target"]).classes(
                                            "text-[13px] text-slate-800 font-semibold leading-snug"
                                        ).style(
                                            "white-space:normal;word-break:break-word"
                                        )
                                    ui.icon("arrow_forward", size="16px").classes(
                                        "text-slate-300 mt-1 shrink-0"
                                    )
                    if not has:
                        ui.label("No intelligence for this segment.").classes(
                            "text-xs text-slate-400 italic"
                        )
        except RuntimeError:
            return

    # ------------------------------------------------------------------
    # Active segment renderer
    # ------------------------------------------------------------------
    def _render_active(seg: dict):
        refs = {"ti": None, "suggestion_bar": None}

        card = ui.card().classes(
            "w-full bg-white active-card p-0 rounded-2xl flex flex-col gap-0 my-6 overflow-hidden"
        )
        seg["_card_id"] = card.id

        with card:
            # --- Header ---
            with ui.row().classes(
                "w-full px-6 py-3 bg-slate-50/50 border-b border-slate-100 justify-between items-center"
            ):
                with ui.row().classes("items-center gap-2"):
                    ui.icon("tag", size="14px", color="slate-400")
                    ui.label(f"SEGMENT {seg['id'] + 1}").classes(
                        "text-[10px] font-black text-slate-400 tracking-[.2em]"
                    )
                with ui.row().classes("items-center gap-1"):
                    if seg["status"] == "done":
                        ui.badge("CONFIRMED", color="emerald-500").classes(
                            "text-[9px] font-bold px-2 py-0.5 rounded-full"
                        )
                    else:
                        ui.badge("DRAFTING", color="blue-400").classes(
                            "text-[9px] font-bold px-2 py-0.5 rounded-full"
                        )

            # --- Source ---
            with ui.column().classes("w-full px-6 pt-6 pb-4 gap-2"):
                ui.label("SOURCE").classes("intel-header")
                with ui.element("div").classes("source-area"):
                    ui.html(_get_highlighted_source(seg["source"])).classes(
                        "text-lg text-slate-800 font-serif leading-relaxed"
                    )

            # --- Target Editor ---
            with ui.column().classes("w-full px-6 py-4 gap-2"):
                ui.label("TARGET").classes("intel-header")
                qa_container = ui.column().classes("w-full gap-1 mb-2")

                ti = (
                    ui.textarea(value=seg["target"])
                    .bind_value(seg, "target")
                    .classes("w-full")
                    .props(
                        "outlined autogrow "
                        'input-style="font-family: Inter, -apple-system, sans-serif; font-size: 16px; line-height: 1.6;"'
                    )
                )
                refs["ti"] = ti

                # =====================================================================
                # INLINE SUGGESTION BAR — cursor-aware
                # =====================================================================
                suggestion_bar = ui.row().classes(
                    "w-full gap-2 items-center min-h-[32px] flex-wrap"
                )
                refs["suggestion_bar"] = suggestion_bar
                _seg_client = context.client

                # Cursor info - stored in JS global variable, read when needed
                _cursor_info = {"start": 0, "end": 0, "word": ""}

                # Skip suggestions on initial value set - only show after user types
                _initialized = False

                # JS to track cursor position in REAL-TIME on client side
                # Stores result in window.__cursor_info_{ti.id} global variable
                _CURSOR_TRACK_JS = f"""
                    (function() {{
                        const el = getHtmlElement({ti.id});
                        if (!el) return;
                        const ta = el.querySelector('textarea');
                        if (!ta) return;

                        function updateCursor() {{
                            const v = ta.value;
                            const p = ta.selectionStart;
                            if (typeof p !== 'number' || p < 0 || p > v.length) return;
                            const re = /[a-zA-Z0-9_čšžČŠŽà-ž]/i;
                            let s = p, e = p;
                            while (s > 0 && re.test(v.charAt(s-1))) s--;
                            while (e < v.length && re.test(v.charAt(e))) e++;
                            const word = v.slice(s, e).trim();
                            window.__cursor_info_{ti.id} = {{word: word || "", start: s, end: e}};
                        }}

                        ta.addEventListener('click', updateCursor);
                        ta.addEventListener('input', updateCursor);  // 'input' fires after value updates, 'keyup' fires before
                        ta.addEventListener('select', updateCursor);
                        // Initialize
                        updateCursor();
                    }})()
                """
                # _CURSOR_TRACK_JS injected in _bg_init after client connects.
                # Injecting here (render time) fails — client not yet connected.

                # JS to read cursor info from global variable
                _GET_CURSOR_JS = f"""
                    (function() {{
                        return window.__cursor_info_{ti.id} || {{word: "", start: 0, end: 0}};
                    }})()
                """

                async def _suggestions_build():
                    """Build suggestions for word at cursor position."""
                    if suggestion_bar.is_deleted or ti.is_deleted:
                        return

                    # await to receive return value — without await, AwaitableResponse
                    # fires-and-forgets and info is never a dict.
                    # Use _seg_client: context.client unavailable in background task.
                    try:
                        info = await _seg_client.run_javascript(
                            _GET_CURSOR_JS, timeout=5.0
                        )
                        if info and isinstance(info, dict):
                            current_word = info.get("word", "")
                            _cursor_info.update(
                                start=info.get("start", 0),
                                end=info.get("end", 0),
                                word=current_word,
                            )
                        else:
                            current_word = ""
                    except Exception:
                        current_word = ""

                    if not current_word or len(current_word) < 2:
                        with _seg_client:
                            if not suggestion_bar.is_deleted:
                                suggestion_bar.clear()
                        return

                    src_l, tgt_l = ws["lang_pair"].split("->")
                    loop = asyncio.get_running_loop()

                    def _query_glossary():
                        if not glossary:
                            return []
                        try:
                            hits = glossary.lookup_terms(current_word, src_l, tgt_l)
                            result = [
                                h["target_term"]
                                for h in hits
                                if h["target_term"]
                                and h["target_term"]
                                .lower()
                                .startswith(current_word.lower())
                            ]
                            return result
                        except Exception:
                            return []

                    def _query_kg():
                        if not kg:
                            return []
                        try:
                            hints = kg.get_inline_hints(
                                word_prefix=current_word[:3],
                                source_text=seg["source"],
                                source_lang=src_l,
                                target_lang=tgt_l,
                                max_hints=5,
                            )
                            result = [
                                h.get("term", "")
                                for h in hints
                                if h.get("term")
                                and h.get("term", "")
                                .lower()
                                .startswith(current_word.lower())
                            ]
                            return result
                        except Exception:
                            return []

                    def _query_tm():
                        if not tm:
                            return []
                        try:
                            hits = tm.search_prefix(current_word)[:5]
                            result = [
                                " ".join(t.split()[:2])
                                for t in hits
                                if t and t.lower().startswith(current_word.lower())
                            ]
                            return result
                        except Exception:
                            return []

                    try:
                        glos_results, kg_results, tm_results = await asyncio.gather(
                            loop.run_in_executor(None, _query_glossary),
                            loop.run_in_executor(None, _query_kg),
                            loop.run_in_executor(None, _query_tm),
                        )
                    except Exception as e:
                        print(f"[Inline] query error: {e}")
                        return

                    # Merge and dedup
                    all_suggestions = list(
                        dict.fromkeys(glos_results + kg_results + tm_results)
                    )

                    # All UI mutations from background task need 'with _seg_client:'
                    with _seg_client:
                        if suggestion_bar.is_deleted:
                            return
                        suggestion_bar.clear()
                        if not all_suggestions:
                            return
                        with suggestion_bar:
                            ui.label("SUGGEST:").classes(
                                "text-[9px] font-black text-slate-400 tracking-widest self-center mr-1"
                            )
                            for idx, sugg in enumerate(all_suggestions[:3]):
                                kind = (
                                    "glossary"
                                    if sugg in glos_results
                                    else "kg"
                                    if sugg in kg_results
                                    else "tm"
                                )
                                color = (
                                    "bg-emerald-100 text-emerald-800"
                                    if kind == "glossary"
                                    else "bg-blue-100 text-blue-800"
                                    if kind == "kg"
                                    else "bg-slate-100 text-slate-600"
                                )
                                text = sugg + (" ↹" if idx == 0 else "")
                                st, en = _cursor_info["start"], _cursor_info["end"]
    
                                def _chip_click(s=sugg, st=st, en=en):
                                    js = f"""
                                        (function() {{
                                            const el = getHtmlElement({ti.id});
                                            if (!el) return;
                                            const ta = el.querySelector('textarea');
                                            if (!ta) return;
                                            ta.setRangeText({json.dumps(s)}, {st}, {en}, 'end');
                                            ta.dispatchEvent(new Event('input', {{bubbles: true}}));
                                            ta.focus();
                                        }})()
                                    """
                                    try:
                                        context.client.run_javascript(js)
                                    except Exception:
                                        pass
    
                                ui.chip(text, on_click=_chip_click).props(
                                    "dense clickable"
                                ).classes(f"{color} text-[11px] font-bold px-2")
                    # end with _seg_client

                # --- Typing handler - simple debounce with direct call ---
                _debounce_task = None

                async def _typing_handler(e: events.ValueChangeEventArguments):
                    nonlocal _debounce_task
                    if not _initialized:
                        return  # Skip on initial value set
                    if _debounce_task and not _debounce_task.done():
                        _debounce_task.cancel()
                    _debounce_task = background_tasks.create(
                        _debounced_suggestions(), name="suggestions"
                    )

                async def _debounced_suggestions():
                    await asyncio.sleep(0.1)
                    try:
                        await _suggestions_build()
                    except Exception:
                        pass

                # Mark as initialized before handler is registered
                _initialized = True
                ti.on_value_change(_typing_handler)

                # --- QA ---
                def _qa_render():
                    if qa_container.is_deleted:
                        return
                    warnings = _run_qa(seg, ws["lang_pair"])
                    qa_container.clear()
                    with qa_container:
                        for w in warnings:
                            color = (
                                "bg-red-50 text-red-700 border-red-100"
                                if w["type"] == "error"
                                else "bg-amber-50 text-amber-700 border-amber-100"
                            )
                            icon = "error" if w["type"] == "error" else "warning"
                            with ui.row().classes(
                                f"w-full {color} px-3 py-2 rounded-lg text-xs items-center gap-2 border"
                            ):
                                ui.icon(icon, size="16px")
                                ui.label(w["message"]).classes("font-medium")

                ti.on_value_change(lambda e: _qa_render())

                # --- Keyboard shortcuts ---
                # Uses js_handler so Tab preventDefault works correctly in NiceGUI 3.x.
                # - js_handler runs first on the client; calls emit(e) to forward to Python.
                # - e.event.preventDefault() blocks browser's native Tab focus-jump,
                #   but only when a ghost suggestion is visible (avoids breaking normal Tab).
                # - The old window.dispatchEvent re-emit pattern is NOT used here because
                #   it causes recursion in NiceGUI 3.1.0+.
                async def _kbd_handler(e: events.KeyEventArguments):
                    """Handle keyboard events using verified NiceGUI APIs."""
                    if (
                        e.key.enter
                        and (e.modifiers.ctrl or e.modifiers.meta)
                        and e.action.keydown
                    ):
                        await _confirm_segment()
                        return

                    # Tab completes first autocomplete suggestion at cursor
                    if e.key.tab and e.action.keydown:
                        # await to receive return value; use _seg_client
                        try:
                            info = await _seg_client.run_javascript(
                                _GET_CURSOR_JS, timeout=5.0
                            )
                            if info and isinstance(info, dict):
                                _cursor_info.update(
                                    start=info.get("start", 0),
                                    end=info.get("end", 0),
                                    word=info.get("word", ""),
                                )
                        except Exception:
                            pass
                        current_word = _cursor_info.get("word", "")
                        start, end = _cursor_info["start"], _cursor_info["end"]
                        if not current_word or len(current_word) < 2:
                            return

                        src_l, tgt_l = ws["lang_pair"].split("->")
                        replacement = None

                        # Try KG first, then glossary
                        if kg:
                            try:
                                hints = kg.get_inline_hints(
                                    word_prefix=current_word[:3],
                                    source_text=seg["source"],
                                    source_lang=src_l,
                                    target_lang=tgt_l,
                                    max_hints=1,
                                )
                                if hints:
                                    replacement = hints[0].get("term")
                            except Exception as ex:
                                print(f"[Inline] tab-KG error: {ex}")

                        if not replacement and glossary:
                            try:
                                hits = glossary.lookup_terms(current_word, src_l, tgt_l)
                                if hits:
                                    replacement = hits[0]["target_term"]
                            except Exception as ex:
                                print(f"[Inline] tab-glossary error: {ex}")

                        if replacement:
                            try:
                                js = f"""
                                    (function() {{
                                        const el = getHtmlElement({ti.id});
                                        if (!el) return;
                                        const ta = el.querySelector('textarea');
                                        if (!ta) return;
                                        ta.setRangeText({json.dumps(replacement)}, {start}, {end}, 'end');
                                        ta.dispatchEvent(new Event('input', {{bubbles: true}}));
                                        ta.focus();
                                    }})()
                                """
                                await _seg_client.run_javascript(js, timeout=5.0)
                                with _seg_client:
                                    if not suggestion_bar.is_deleted:
                                        suggestion_bar.clear()
                            except Exception as ex:
                                print(f"[Inline] tab-insert error: {ex}")

                # Keyboard handler using verified NiceGUI 3.11.1 API
                # ui.keyboard with on_key parameter receives proper KeyEventArguments
                ui.keyboard(on_key=_kbd_handler, ignore=[])

            # --- Static inline panel (TM / Glossary / Concordance) ---
            static_panel = ui.column().classes(
                "w-full px-6 py-4 bg-slate-50/30 border-t border-slate-100 gap-3"
            )

            # --- Footer actions ---
            with ui.row().classes(
                "w-full px-6 py-4 bg-slate-50/50 border-t border-slate-100 justify-between items-center"
            ):
                with ui.row().classes("items-center gap-2"):
                    ui.button(
                        icon="auto_awesome",
                        on_click=lambda: background_tasks.create(
                            _ai(seg, refs["ti"], force=True), name="ai_regenerate"
                        ),
                    ).props("flat round dense size=md color=accent").tooltip(
                        "Regenerate AI Draft"
                    )
                    ui.separator().props("vertical").classes("mx-2 h-6 opacity-20")
                    if seg["id"] > 0:
                        ui.button(
                            icon="keyboard_arrow_up",
                            on_click=lambda: _set_active(seg["id"] - 1),
                        ).props("flat round dense size=md color=slate-400")
                    if seg["id"] < len(ws["segments"]) - 1:
                        ui.button(
                            icon="keyboard_arrow_down",
                            on_click=lambda: _set_active(seg["id"] + 1),
                        ).props("flat round dense size=md color=slate-400")

                with ui.row().classes("items-center gap-4"):
                    ui.label("TAB replaces word • ⌘ ENTER confirms").classes(
                        "text-[10px] text-slate-400 font-bold uppercase tracking-wider"
                    )
                    ui.button("CONFIRM", on_click=lambda: _confirm_segment()).props(
                        "unelevated rounded color=emerald-500"
                    ).classes(
                        "px-8 py-2 font-black tracking-[.2em] text-[11px] shadow-lg shadow-emerald-200/50"
                    )

            # --- Background init ---
            async def _bg_init():
                await asyncio.sleep(0.1)
                if ws["project_id"] not in GLOBAL_VOCAB:
                    GLOBAL_VOCAB[ws["project_id"]] = set()
                v = GLOBAL_VOCAB[ws["project_id"]]
                for s in ws["segments"]:
                    for t in [s["source"], s["target"]]:
                        if t:
                            v.update(
                                re.findall(
                                    r"[\w\u010D\u0161\u017E\u010C\u0160\u017D]{3,}", t
                                )
                            )

                try:
                    # Inject cursor tracker now — client is connected after the sleep above.
                    # Must be awaited so it runs before the user starts typing.
                    # fire-and-forget (no await) at render time silently fails because
                    # the WebSocket handshake hasn't completed yet.
                    await _seg_client.run_javascript(_CURSOR_TRACK_JS, timeout=5.0)

                    if not seg["target"].strip():
                        background_tasks.create(_ai(seg, refs["ti"]), name="ai_init")
                    if not static_panel.is_deleted:
                        background_tasks.create(
                            _inline_panel(seg, refs["ti"], static_panel),
                            name="inline_panel",
                        )
                except Exception as e:
                    print(f"[Inline] BG init error: {e}")

            background_tasks.create(_bg_init(), name="bg_init")

    def _render_inactive(seg: dict):
        border = (
            "border-emerald-400" if seg["status"] == "done" else "border-transparent"
        )
        bg = "bg-white" if seg["status"] == "done" else "bg-slate-50/30"
        qa = _run_qa(seg, ws["lang_pair"]) if seg["target"].strip() else []
        has_error = any(w["type"] == "error" for w in qa)
        has_warn = any(w["type"] == "warning" for w in qa)

        with (
            ui.row()
            .classes(
                f"w-full {bg} border-l-4 {border} px-6 py-4 rounded-xl cursor-pointer seg-row no-wrap gap-8 items-center border border-slate-100 mb-2"
            )
            .on("click", lambda s=seg: _set_active(s["id"]))
        ):
            with ui.column().classes("w-1/2 gap-1"):
                ui.label(f"#{seg['id'] + 1}").classes(
                    "text-[8px] font-black text-slate-300 tracking-widest"
                )
                ui.label(seg["source"]).classes(
                    "text-[14px] text-slate-600 font-serif leading-relaxed line-clamp-2"
                )

            with ui.row().classes("w-1/2 items-center gap-3 no-wrap"):
                lbl = (
                    seg["target"]
                    if seg["target"].strip()
                    else "Waiting for translation..."
                )
                sty = (
                    "text-slate-800 font-medium"
                    if seg["target"].strip()
                    else "text-slate-300 italic font-light"
                )
                ui.label(lbl).classes(
                    f"text-[14px] {sty} leading-relaxed line-clamp-2 flex-1"
                )
                if has_error:
                    ui.icon("error", color="negative", size="18px").tooltip(
                        "Critical QA Error"
                    )
                elif has_warn:
                    ui.icon("warning", color="warning", size="18px").tooltip(
                        "QA Warning"
                    )
                if seg["status"] == "done":
                    ui.icon("check_circle", color="emerald-400", size="18px")

    async def _batch():
        if ws["is_batch"]:
            return
        ws["is_batch"] = True
        cli = context.client

        with cli:
            _bar.refresh()

        src, tgt = ws["lang_pair"].split("->")
        loop = asyncio.get_running_loop()

        for i, seg in enumerate(ws["segments"]):
            if not ws["is_batch"]:
                break
            if seg["status"] == "done" or seg["target"].strip():
                continue
            try:
                a = (
                    tm.lookup_fuzzy(seg["source"], threshold=90.0, limit=1)
                    if tm
                    else []
                )
                b = glossary.lookup_terms(seg["source"], src, tgt) if glossary else []
                c = tm.search_concordance(seg["source"], top_n=2) if tm else []
                _, text = await loop.run_in_executor(
                    llm_executor,
                    lambda s=seg, a=a, b=b, c=c: translator.translate(
                        s["source"], src, tgt, a, b, c
                    ),
                )
                seg["target"] = text
                if i % 3 == 0:
                    _autosave()
                    with cli:
                        _flow.refresh()
                        _bar.refresh()
            except Exception as ex:
                print(f"[Batch] {ex}")

        with cli:
            ws["is_batch"] = False
            _autosave()
            _flow.refresh()
            _bar.refresh()
            ui.notify("Batch complete!", type="positive")

    def _stop_batch():
        ws["is_batch"] = False
        _bar.refresh()
        ui.notify("Batch stopped", type="warning")

    def _export():
        docx_path = PROJECTS_DIR / f"{project_id}.docx"
        if not docx_path.exists():
            return ui.notify("Original DOCX not found", type="negative")
        doc = docx.Document(str(docx_path))
        by_src = {s["source"]: s for s in ws["segments"]}
        for para in doc.paragraphs:
            key = para.text.strip()
            if key in by_src and by_src[key]["target"].strip():
                para.text = by_src[key]["target"]
        stream = io.BytesIO()
        doc.save(stream)
        ui.download(stream.getvalue(), f"translated_{ws['filename']}")

    def _export_txt():
        lines = []
        for s in ws["segments"]:
            if s["target"].strip():
                lines.append(s["target"].strip())
            else:
                lines.append(s["source"].strip())
        content = "\n\n".join(lines)
        ui.download(content.encode("utf-8"), f"translated_{ws['filename']}.txt")

    @ui.refreshable
    def _bar():
        with ui.row().classes(
            "w-full items-center justify-between bg-white border-b border-slate-200 px-5 py-3 shadow-sm"
        ):
            with ui.row().classes("items-center gap-3"):
                ui.button(
                    icon="arrow_back", on_click=lambda: ui.navigate.to("/")
                ).props("flat round dense").tooltip("Back to projects")
                with ui.column().classes("gap-0"):
                    ui.label(ws["filename"]).classes(
                        "text-sm font-bold text-slate-700 leading-none"
                    )
                    ui.label(ws["lang_pair"]).classes("text-[10px] text-slate-400")
                ui.button(icon="menu_book", on_click=_open_glossary).props(
                    "flat round dense size=sm color=slate-400"
                ).tooltip("Add Glossary Term")

            pct = _progress()
            with ui.column().classes("w-44 items-center gap-0.5"):
                ui.label(f"{int(pct * 100)}%").classes(
                    "text-[10px] font-bold text-slate-500 uppercase tracking-widest"
                )
                ui.linear_progress(value=pct, color="positive").props(
                    'size="6px" :show-value="false" track-color="slate-200"'
                ).classes("w-full rounded-full")

            with ui.row().classes("gap-2"):
                if ws["is_batch"]:
                    ui.button(
                        "Stop Translating", icon="stop", on_click=_stop_batch
                    ).props("outline rounded dense color=negative")
                else:
                    ui.button(
                        "Auto-translate", icon="auto_awesome", on_click=_batch
                    ).props("outline rounded dense color=accent")
                with ui.button("Export", icon="file_download").props(
                    "rounded unelevated dense color=positive"
                ):
                    with ui.menu():
                        ui.menu_item("Export as .docx", on_click=_export)
                        ui.menu_item("Export as .txt", on_click=_export_txt)

    @ui.refreshable
    def _flow():
        for idx, seg in enumerate(ws["segments"]):
            if idx == ws["active_index"]:
                _render_active(seg)
            else:
                _render_inactive(seg)

    with ui.column().classes("w-full h-screen bg-slate-50 no-wrap"):
        _bar()
        with ui.column().classes("w-full flex-1 overflow-y-auto pb-24 pt-4"):
            with ui.column().classes("w-full max-w-4xl mx-auto px-4 gap-1"):
                _flow()

    ui.timer(1.0, _scan_tm_for_kg, once=True)


def _apply_colors():
    ui.colors(
        primary="#0f172a",
        secondary="#334155",
        positive="#10b981",
        accent="#3b82f6",
        negative="#ef4444",
    )


if __name__ in {"__main__", "__mp_main__"}:
    ui.run(title="Zen Translator", favicon="✨", port=8080, show=True)
